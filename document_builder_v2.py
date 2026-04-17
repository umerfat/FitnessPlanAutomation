"""
Build a professional, colorful Word document from the structured fitness plan.
Brand: UMER HURRAH — TRANSFORMATION PROTOCOL
"""

import os

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

import config


# Brand color palette
BRAND_PRIMARY = RGBColor(0x6C, 0x28, 0xD2)     # Purple
BRAND_DARK = RGBColor(0x2D, 0x0A, 0x6E)         # Deep purple
BRAND_ACCENT = RGBColor(0xFF, 0x6B, 0x35)        # Orange accent
BRAND_GREEN = RGBColor(0x27, 0xAE, 0x60)         # Green for alternatives
BRAND_BLUE = RGBColor(0x29, 0x80, 0xB9)          # Blue for info
BRAND_RED = RGBColor(0xE7, 0x4C, 0x3C)           # Red for warnings
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK_TEXT = RGBColor(0x2C, 0x2C, 0x2C)
LIGHT_GRAY = RGBColor(0x95, 0xA5, 0xA6)
BG_LIGHT_PURPLE = "E8D5F5"
BG_LIGHT_ORANGE = "FFF3ED"
BG_LIGHT_GREEN = "E8F8F0"
BG_LIGHT_BLUE = "EBF5FB"


def build_document(client_data: dict, metrics: dict, plan: dict) -> str:
    doc = Document()
    _set_default_font(doc)
    _set_narrow_margins(doc)

    name = client_data.get("Full Name", "Client")
    age = metrics["age"]
    gender = metrics["gender"]
    goal = client_data.get("Primary Fitness Goal", "")
    cal = metrics["calories"]
    macros = metrics["macros"]

    # === BRANDED HEADER ===
    _add_branded_header(doc)

    # === CLIENT SUMMARY with body stats table ===
    _add_section_heading(doc, "📋  CLIENT SUMMARY", BRAND_PRIMARY)
    _add_client_summary_table(doc, client_data, metrics, plan)
    doc.add_paragraph()

    # === NUTRITION PLAN ===
    _add_section_heading(doc, "🍽️  NUTRITION PLAN", BRAND_ACCENT)
    _add_macro_summary_bar(doc, macros, cal)
    nutrition = plan.get("nutrition_plan", {})
    meals = nutrition.get("meals", [])
    for i, meal in enumerate(meals, 1):
        _add_meal_card(doc, i, meal)

    # Nutrition rules
    rules = nutrition.get("rules", [])
    if rules:
        _add_rules_box(doc, rules)
    doc.add_paragraph()

    # === TRAINING SPLIT ===
    _add_section_heading(doc, "💪  TRAINING SPLIT", BRAND_BLUE)
    training = plan.get("training_split", {})
    days = training.get("days", [])
    for day_plan in days:
        _add_workout_day(doc, day_plan)

    # Weekend rest
    rest_para = doc.add_paragraph()
    rest_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rest_para.paragraph_format.space_before = Pt(12)
    run = rest_para.add_run("🧘 Saturday & Sunday: Rest / Active Recovery (light walk, stretching, yoga)")
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = BRAND_PRIMARY
    doc.add_paragraph()

    # === SUPPLEMENTS ===
    supplements = plan.get("supplements", [])
    if supplements:
        _add_section_heading(doc, "💊  RECOMMENDED SUPPLEMENTS", BRAND_GREEN)
        _add_supplements_table(doc, supplements)
        doc.add_paragraph()

    # === LIFESTYLE UPGRADE ===
    lifestyle = plan.get("lifestyle_upgrade", [])
    if lifestyle:
        _add_section_heading(doc, "⚡  LIFESTYLE UPGRADE", BRAND_DARK)
        _add_lifestyle_tips(doc, lifestyle)

    # === FOOTER ===
    _add_footer(doc)

    # Save
    safe_name = name.replace(" ", "_").replace("/", "_")
    filename = f"{safe_name}_plan.docx"
    filepath = os.path.join(config.OUTPUT_DIR, filename)
    os.makedirs(config.OUTPUT_DIR, exist_ok=True)
    doc.save(filepath)
    return filepath


# ─── BRANDED HEADER ───────────────────────────────────────


def _add_branded_header(doc):
    """Big colorful brand header at the top."""
    # Brand name
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(config.BRAND_NAME.upper())
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = BRAND_PRIMARY

    # Subtitle
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(6)
    run = sub.add_run("TRANSFORMATION PROTOCOL")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = BRAND_ACCENT

    # Colored divider line via table
    _add_colored_divider(doc, BRAND_PRIMARY)
    doc.add_paragraph()


# ─── CLIENT SUMMARY ───────────────────────────────────────


def _add_client_summary_table(doc, client_data, metrics, plan):
    """Colored table with body stats."""
    name = client_data.get("Full Name", "Client")
    age = metrics["age"]
    gender = metrics["gender"]
    height = metrics["height_cm"]
    weight = metrics["weight_kg"]
    bmi = metrics["bmi"]
    cal = metrics["calories"]
    macros = metrics["macros"]
    body_fat = metrics["body_fat_pct"]
    bf_cat = metrics["body_fat_category"]
    lean_mass = metrics["lean_mass_kg"]
    fat_mass = metrics["fat_mass_kg"]
    whr = metrics.get("whr")
    hydration = metrics["hydration_liters"]

    # Info table - 2 columns
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    stats = [
        ("Client", f"{name}  |  Age: {age}  |  {gender}"),
        ("Height / Weight", f"{height} cm  |  {weight} kg"),
        ("BMI", f"{bmi['bmi']}  ({bmi['category']})"),
        ("Est. Body Fat", f"{body_fat}%  ({bf_cat})"),
        ("Lean Body Mass", f"{lean_mass} kg"),
        ("Fat Mass", f"{fat_mass} kg"),
        ("Goal", client_data.get("Primary Fitness Goal", "")),
        ("Daily Target", f"~{cal['target_calories']} kcal  ({cal['strategy']})"),
        ("Macro Split", f"P: {macros['protein_g']}g  |  C: {macros['carbs_g']}g  |  F: {macros['fats_g']}g"),
        ("Hydration Target", f"{hydration}L / day"),
    ]

    if whr:
        stats.insert(6, ("Waist-to-Hip Ratio", f"{whr['whr']}  ({whr['risk']})"))

    for label, value in stats:
        row = table.add_row()
        # Label cell
        cell_label = row.cells[0]
        cell_label.width = Cm(5)
        p = cell_label.paragraphs[0]
        run = p.add_run(label)
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = BRAND_PRIMARY
        _shade_cell(cell_label, BG_LIGHT_PURPLE)
        cell_label.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Value cell
        cell_val = row.cells[1]
        p = cell_val.paragraphs[0]
        run = p.add_run(value)
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_TEXT
        cell_val.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    _style_table_borders(table)

    # Physique insight
    insight = plan.get("physique_insight", "")
    if insight:
        doc.add_paragraph()
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(4)
        run = para.add_run("🔍 Physique Insight: ")
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = BRAND_DARK
        run = para.add_run(insight)
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = DARK_TEXT


# ─── NUTRITION ────────────────────────────────────────────


def _add_macro_summary_bar(doc, macros, cal):
    """Colored macro summary."""
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    items = [
        ("🔥 Calories", f"{cal['target_calories']} kcal", BG_LIGHT_ORANGE),
        ("🥩 Protein", f"{macros['protein_g']}g", BG_LIGHT_PURPLE),
        ("🍚 Carbs", f"{macros['carbs_g']}g", BG_LIGHT_BLUE),
        ("🥑 Fats", f"{macros['fats_g']}g", BG_LIGHT_GREEN),
    ]

    for i, (label, value, bg) in enumerate(items):
        cell = table.rows[0].cells[i]
        _shade_cell(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}\n")
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = DARK_TEXT
        run = p.add_run(value)
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = BRAND_PRIMARY

    _style_table_borders(table)
    doc.add_paragraph()


def _add_meal_card(doc, num, meal):
    """A visually distinct card for each meal with alternative."""
    timing = meal.get("timing", f"Meal {num}")
    time_sug = meal.get("time_suggestion", "")
    items = meal.get("items", "")
    alt = meal.get("alternative", "")
    p_val = meal.get("protein", 0)
    c_val = meal.get("carbs", 0)
    f_val = meal.get("fats", 0)

    # Meal header
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(f"  Meal {num}  ")
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = WHITE
    _highlight_run(run, "6C28D2")

    time_label = f"  {timing}" + (f" ({time_sug})" if time_sug else "")
    run = para.add_run(time_label)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = BRAND_DARK

    # Primary meal
    p_items = doc.add_paragraph()
    p_items.paragraph_format.left_indent = Inches(0.3)
    p_items.paragraph_format.space_before = Pt(2)
    p_items.paragraph_format.space_after = Pt(2)
    run = p_items.add_run("▸ ")
    run.font.color.rgb = BRAND_ACCENT
    run.font.size = Pt(10)
    run = p_items.add_run(items)
    run.font.size = Pt(10)
    run.font.color.rgb = DARK_TEXT

    # Alternative
    if alt:
        p_alt = doc.add_paragraph()
        p_alt.paragraph_format.left_indent = Inches(0.3)
        p_alt.paragraph_format.space_before = Pt(0)
        p_alt.paragraph_format.space_after = Pt(2)
        run = p_alt.add_run("↻ Alternative: ")
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = BRAND_GREEN
        run = p_alt.add_run(alt)
        run.font.size = Pt(9)
        run.font.color.rgb = BRAND_GREEN

    # Macro line
    p_macro = doc.add_paragraph()
    p_macro.paragraph_format.left_indent = Inches(0.3)
    p_macro.paragraph_format.space_before = Pt(0)
    p_macro.paragraph_format.space_after = Pt(4)
    run = p_macro.add_run(f"P: {p_val}g  |  C: {c_val}g  |  F: {f_val}g")
    run.font.size = Pt(9)
    run.font.color.rgb = LIGHT_GRAY
    run.font.italic = True


def _add_rules_box(doc, rules):
    """Colored box with nutrition rules."""
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    _shade_cell(cell, BG_LIGHT_ORANGE)

    p = cell.paragraphs[0]
    run = p.add_run("📌 Nutrition Rules")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = BRAND_ACCENT

    for i, rule in enumerate(rules, 1):
        p = cell.add_paragraph()
        run = p.add_run(f"  {i}. {rule}")
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_TEXT

    _style_table_borders(table)


# ─── TRAINING ─────────────────────────────────────────────


def _add_workout_day(doc, day_plan):
    """Colored workout day with exercises and alternatives."""
    day = day_plan.get("day", "")
    focus = day_plan.get("focus", "")
    exercises = day_plan.get("exercises", [])

    # Day header with colored background
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    _shade_cell(cell, BG_LIGHT_BLUE)
    p = cell.paragraphs[0]
    run = p.add_run(f"  {day}  —  {focus}")
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = BRAND_BLUE
    _style_table_borders(table)

    # Exercise table
    ex_table = doc.add_table(rows=1, cols=4)
    ex_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    headers = ["#", "Exercise", "Sets × Reps", "Rest"]
    for i, h in enumerate(headers):
        cell = ex_table.rows[0].cells[i]
        _shade_cell(cell, "2980B9")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = WHITE

    for j, ex in enumerate(exercises, 1):
        row = ex_table.add_row()
        ex_name = ex.get("name", "")
        alt = ex.get("alternative", "")
        sets_reps = ex.get("sets_reps", "")
        rest = ex.get("rest", "")
        notes = ex.get("notes", "")

        # Number
        p = row.cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(j))
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = BRAND_BLUE

        # Exercise name + alternative
        p = row.cells[1].paragraphs[0]
        run = p.add_run(ex_name)
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = DARK_TEXT
        if alt:
            p2 = row.cells[1].add_paragraph()
            run = p2.add_run(f"↻ Alt: {alt}")
            run.font.size = Pt(8)
            run.font.color.rgb = BRAND_GREEN
            run.font.italic = True
        if notes:
            p3 = row.cells[1].add_paragraph()
            run = p3.add_run(f"💡 {notes}")
            run.font.size = Pt(8)
            run.font.color.rgb = LIGHT_GRAY

        # Sets/Reps
        p = row.cells[2].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(sets_reps)
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_TEXT

        # Rest
        p = row.cells[3].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(rest)
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_TEXT

        # Alternate row shading
        if j % 2 == 0:
            for cell in row.cells:
                _shade_cell(cell, "F8F9FA")

    _style_table_borders(ex_table)
    doc.add_paragraph()


# ─── SUPPLEMENTS ──────────────────────────────────────────


def _add_supplements_table(doc, supplements):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, h in enumerate(["#", "Supplement & Dosage"]):
        cell = table.rows[0].cells[i]
        _shade_cell(cell, "27AE60")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(h)
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = WHITE

    for i, supp in enumerate(supplements, 1):
        row = table.add_row()
        p = row.cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(i))
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = BRAND_GREEN

        p = row.cells[1].paragraphs[0]
        run = p.add_run(supp)
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_TEXT

        if i % 2 == 0:
            for cell in row.cells:
                _shade_cell(cell, "F8F9FA")

    table.rows[0].cells[0].width = Cm(1.5)
    _style_table_borders(table)


# ─── LIFESTYLE ────────────────────────────────────────────


def _add_lifestyle_tips(doc, tips):
    icons = ["🌙", "🚶", "📱", "🧠", "💧", "🎯", "⏰", "🥗"]
    for i, tip in enumerate(tips):
        icon = icons[i % len(icons)]
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.3)
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after = Pt(4)
        run = para.add_run(f"  {icon}  ")
        run.font.size = Pt(11)
        run = para.add_run(tip)
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_TEXT


# ─── FOOTER ───────────────────────────────────────────────


def _add_footer(doc):
    doc.add_paragraph()
    _add_colored_divider(doc, BRAND_PRIMARY)
    doc.add_paragraph()

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f"— {config.BRAND_NAME.upper()} —")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = BRAND_PRIMARY

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run("Your transformation starts now. Stay consistent. Trust the process.")
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = BRAND_ACCENT

    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disclaimer.add_run(
        "This plan is personalized based on the information you provided. "
        "Consult a medical professional before beginning any new fitness or nutrition program."
    )
    run.font.size = Pt(8)
    run.font.color.rgb = LIGHT_GRAY


# ─── HELPERS ──────────────────────────────────────────────


def _set_default_font(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = DARK_TEXT


def _set_narrow_margins(doc):
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)


def _add_section_heading(doc, text, color):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(20)
    para.paragraph_format.space_after = Pt(10)
    run = para.add_run(text)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = color


def _add_colored_divider(doc, color):
    """Add a thin colored line as a visual divider."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    hex_color = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}" if isinstance(color, RGBColor) else str(color)
    _shade_cell(cell, hex_color)
    p = cell.paragraphs[0]
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    run = p.add_run("")
    run.font.size = Pt(2)


def _shade_cell(cell, hex_color):
    """Apply a background color to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _highlight_run(run, hex_color):
    """Apply highlight/shading to a run of text."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    run._r.get_or_add_rPr().append(shading)


def _style_table_borders(table):
    """Apply light borders to a table."""
    tbl = table._tbl
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
        '</w:tblBorders>'
    )
    tbl.tblPr.append(borders)
