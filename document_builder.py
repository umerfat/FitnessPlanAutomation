"""
Build a professional Word document from the structured fitness plan.
Matches the 'UMER HURRAH — TRANSFORMATION PROTOCOL' style.
"""

import os
import re

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

import config


# Brand colors
BRAND_PURPLE = RGBColor(0x4A, 0x14, 0x8C)
DARK_TEXT = RGBColor(0x1A, 0x1A, 0x1A)
GRAY_TEXT = RGBColor(0x55, 0x55, 0x55)


def build_document(client_data: dict, metrics: dict, plan: dict) -> str:
    """
    Create a Word document for the client's fitness plan.
    Returns the file path of the generated document.
    """
    doc = Document()
    _set_default_font(doc)

    name = client_data.get("Full Name", "Client")
    age = metrics.get("age", "")
    goal = client_data.get("Primary Fitness Goal", "")
    cal = metrics.get("calories", {})
    macros = metrics.get("macros", {})

    # === TITLE ===
    _add_title(doc, f"{config.BRAND_NAME} — TRANSFORMATION PROTOCOL")

    # === CLIENT SUMMARY ===
    _add_heading(doc, "CLIENT SUMMARY")
    _add_bold_line(doc, "Client:", f"{name} | Age: {age}")
    _add_bold_line(doc, "Goal:", goal)
    _add_bold_line(doc, "Daily Target Calories:", f"~{cal['target_calories']} kcal")
    _add_bold_line(
        doc,
        "Macro Split:",
        f"Protein: {macros['protein_g']}g | Carbs: {macros['carbs_g']}g | Fats: {macros['fats_g']}g",
    )

    # Physique insight
    insight = plan.get("physique_insight", "")
    _add_bold_line(doc, "Physique Insight:", insight)
    doc.add_paragraph()  # spacer

    # === DETAILED NUTRITION PLAN ===
    _add_heading(doc, "DETAILED NUTRITION PLAN (WITH MACROS)")
    nutrition = plan.get("nutrition_plan", {})
    meals = nutrition.get("meals", [])

    for i, meal in enumerate(meals, 1):
        timing = meal.get("timing", f"Meal {i}")
        time_sug = meal.get("time_suggestion", "")
        items = meal.get("items", "")
        p = meal.get("protein", 0)
        c = meal.get("carbs", 0)
        f_val = meal.get("fats", 0)

        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(2)

        # Meal number + timing
        run_num = para.add_run(f"{i}   ")
        run_num.font.size = Pt(11)
        run_num.font.bold = True

        time_label = f"{timing}" + (f" ({time_sug})" if time_sug else "")
        run_timing = para.add_run(f"{time_label}: ")
        run_timing.font.size = Pt(11)
        run_timing.font.bold = True

        run_items = para.add_run(items)
        run_items.font.size = Pt(11)

        # Macro line
        macro_para = doc.add_paragraph()
        macro_para.paragraph_format.space_before = Pt(0)
        macro_para.paragraph_format.space_after = Pt(4)
        macro_para.paragraph_format.left_indent = Inches(0.4)
        run_macro = macro_para.add_run(f"~ Protein: {p}g | Carbs: {c}g | Fats: {f_val}g")
        run_macro.font.size = Pt(10)
        run_macro.font.color.rgb = GRAY_TEXT

    # Nutrition rules
    rules = nutrition.get("rules", [])
    if rules:
        doc.add_paragraph()
        rules_para = doc.add_paragraph()
        run_label = rules_para.add_run("Rules: ")
        run_label.font.bold = True
        run_label.font.size = Pt(11)
        run_text = rules_para.add_run("; ".join(rules))
        run_text.font.size = Pt(11)

    doc.add_paragraph()  # spacer

    # === TRAINING SPLIT ===
    _add_heading(doc, "TRAINING SPLIT")
    training = plan.get("training_split", {})
    days = training.get("days", [])

    for day_plan in days:
        day = day_plan.get("day", "")
        focus = day_plan.get("focus", "")
        exercises = day_plan.get("exercises", [])

        # Day header
        day_para = doc.add_paragraph()
        day_para.paragraph_format.space_before = Pt(12)
        day_para.paragraph_format.space_after = Pt(6)
        run_day = day_para.add_run(f"{day} ({focus})")
        run_day.font.bold = True
        run_day.font.italic = True
        run_day.font.size = Pt(12)

        # Exercise list
        for j, ex in enumerate(exercises, 1):
            ex_name = ex.get("name", "")
            sets_reps = ex.get("sets_reps", "")
            rest = ex.get("rest", "")
            notes = ex.get("notes", "")

            ex_para = doc.add_paragraph()
            ex_para.paragraph_format.space_before = Pt(2)
            ex_para.paragraph_format.space_after = Pt(2)
            ex_para.paragraph_format.left_indent = Inches(0.3)

            run_num = ex_para.add_run(f"{j}   ")
            run_num.font.size = Pt(11)
            run_num.font.bold = True

            line = f"{ex_name} – {sets_reps}"
            if rest:
                line += f" (Rest: {rest})"
            run_ex = ex_para.add_run(line)
            run_ex.font.size = Pt(11)

            if notes:
                note_para = doc.add_paragraph()
                note_para.paragraph_format.left_indent = Inches(0.6)
                note_para.paragraph_format.space_before = Pt(0)
                note_para.paragraph_format.space_after = Pt(2)
                run_note = note_para.add_run(f"→ {notes}")
                run_note.font.size = Pt(9)
                run_note.font.color.rgb = GRAY_TEXT

    # Weekend note
    doc.add_paragraph()
    rest_para = doc.add_paragraph()
    run_rest = rest_para.add_run("Saturday & Sunday: Rest / Active Recovery (light walk, stretching)")
    run_rest.font.italic = True
    run_rest.font.size = Pt(11)

    doc.add_paragraph()  # spacer

    # === SUPPLEMENTS ===
    supplements = plan.get("supplements", [])
    if supplements:
        _add_heading(doc, "RECOMMENDED SUPPLEMENTS")
        for i, supp in enumerate(supplements, 1):
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.3)
            run_num = para.add_run(f"{i}   ")
            run_num.font.bold = True
            run_num.font.size = Pt(11)
            run_text = para.add_run(supp)
            run_text.font.size = Pt(11)
        doc.add_paragraph()

    # === LIFESTYLE UPGRADE ===
    lifestyle = plan.get("lifestyle_upgrade", [])
    if lifestyle:
        _add_heading(doc, "LIFESTYLE UPGRADE")
        for i, tip in enumerate(lifestyle, 1):
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.3)
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after = Pt(3)
            run_num = para.add_run(f"{i}   ")
            run_num.font.bold = True
            run_num.font.size = Pt(11)
            run_text = para.add_run(tip)
            run_text.font.size = Pt(11)

    # === FOOTER ===
    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_footer = footer.add_run(f"— {config.BRAND_NAME} —")
    run_footer.font.size = Pt(10)
    run_footer.font.color.rgb = GRAY_TEXT
    run_footer.font.italic = True

    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_disc = disclaimer.add_run(
        "This plan is personalized based on the information you provided. "
        "Consult a medical professional before beginning any new fitness or nutrition program."
    )
    run_disc.font.size = Pt(8)
    run_disc.font.color.rgb = GRAY_TEXT

    # Save
    safe_name = re.sub(r'[<>:"|?*\\/]', '_', name).replace(" ", "_")
    filename = f"{safe_name}_plan.docx"
    filepath = os.path.join(config.OUTPUT_DIR, filename)
    os.makedirs(config.OUTPUT_DIR, exist_ok=True)
    doc.save(filepath)

    return filepath


def _set_default_font(doc: Document):
    """Set default font for the document."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = DARK_TEXT


def _add_title(doc: Document, text: str):
    """Add a centered, bold title."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(24)
    run = para.add_run(text)
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = BRAND_PURPLE


def _add_heading(doc: Document, text: str):
    """Add a section heading."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(18)
    para.paragraph_format.space_after = Pt(10)
    run = para.add_run(text)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = DARK_TEXT


def _add_bold_line(doc: Document, label: str, value: str):
    """Add a line with bold label and normal value."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    run_label = para.add_run(f"{label} ")
    run_label.font.bold = True
    run_label.font.size = Pt(11)
    run_value = para.add_run(value)
    run_value.font.size = Pt(11)
